#!/usr/bin/env python3
"""Generate NozyWallet technical paper (~20 pages) as Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUTPUT = r"c:\Users\User\NozyWallet\docs\NozyWallet_Technical_Paper_v2.docx"


def set_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(3)


def cover_page(doc):
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NozyWallet")
    r.bold = True
    r.font.size = Pt(36)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        "A Privacy-First Orchard Zcash Wallet:\n"
        "Architecture, Cryptography, and Implementation"
    )
    r2.font.size = Pt(15)
    r2.font.name = "Calibri"
    r2.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        "Technical Report — Extended Edition",
        "",
        "LEONINE DAO",
        "Version 2.3.6.5 (Teriyaki Hot)",
        date.today().strftime("%B %d, %Y"),
        "",
        "github.com/LEONINE-DAO/Nozy-wallet",
    ]
    for i, line in enumerate(lines):
        if i > 0:
            meta.add_run("\n")
        run = meta.add_run(line)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
    doc.add_page_break()


def table_of_contents(doc):
    add_heading(doc, "Table of Contents", 1)
    entries = [
        "Abstract",
        "1. Introduction and Motivation",
        "2. Background: Zcash Shielded Pools",
        "3. The Orchard Protocol and NozyWallet",
        "4. Viewing Keys in Zcash and NozyWallet",
        "5. System Architecture",
        "6. Implemented Features and Release History",
        "7. Security, Privacy, and Deterministic Indexing",
        "8. Product Roadmap: Pure Zcash Stack",
        "9. Conclusion",
        "References",
        "Appendix A: Glossary",
    ]
    for e in entries:
        add_body(doc, e, space_after=2)
    doc.add_page_break()


def abstract(doc):
    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "NozyWallet is an Orchard-first Zcash wallet implemented in Rust. The project enforces "
        "shielded transactions as the only user-facing option: transparent t-address sends and "
        "receives are rejected by policy. Development began in September 2025 on the Zebra (zebrad) "
        "full-node stack with lightwalletd for compact block delivery, replacing dependence on legacy "
        "zcashd infrastructure and node-supplied witness RPCs that had blocked reliable shielded sends.",
    )
    add_body(
        doc,
        "This report provides a comprehensive technical overview suitable for developers, operators, "
        "and reviewers. It explains the Orchard shielded pool at the protocol level, describes how "
        "viewing keys (unified full viewing keys, incoming viewing keys, and full viewing keys) "
        "function in Zcash and how NozyWallet uses them for scanning, hardware-wallet pairing, and "
        "planned selective disclosure for business accounting. It documents system architecture, "
        "witness derivation, compact sync (Zeaking), dynamic-fee pilot alignment, multi-surface "
        "delivery, and mainnet hardening through release v2.3.6.5.",
    )
    add_body(
        doc,
        "Production-ready surface: the nozy command-line interface with operator-run Zebrad and "
        "lightwalletd. Desktop, browser extension, api-server companion, and mobile applications "
        "remain under active development.",
    )


def section1_intro(doc):
    add_heading(doc, "1. Introduction and Motivation", 1)

    add_heading(doc, "1.1 The Transparency Problem in Zcash", 2)
    add_body(
        doc,
        "Zcash is often described as a privacy coin, yet the protocol itself supports two parallel "
        "ledger views. Transparent transactions use t-addresses and behave like Bitcoin: amounts, "
        "sender and receiver addresses, and transaction graphs are visible to any observer. Shielded "
        "pools—Sapling and Orchard—use zero-knowledge proofs to hide these fields. Wallets that "
        "default to transparent flows or that make shielded usage optional increase the probability "
        "that users expose financial activity without intending to.",
    )
    add_body(
        doc,
        "NozyWallet was created to remove that ambiguity at the product layer. The wallet implements "
        "Orchard shielded transactions and unified addresses (u1) as the sole supported path for "
        "send and receive. The privacy module (src/privacy.rs) enforces rejection of transparent "
        "addresses in user-facing validation. This is a deliberate product constraint: NozyWallet "
        "does not claim to improve network-level anonymity sets beyond what Orchard cryptography "
        "provides, but it does prevent accidental transparent leakage through the wallet interface.",
    )

    add_heading(doc, "1.2 Project Scope and Non-Goals", 2)
    add_body(
        doc,
        "NozyWallet is a wallet and companion services repository. It is not a Zcash consensus "
        "node, not a mining client, and not a bridge operator. Consensus validation, block production, "
        "and network propagation are delegated to Zebrad. Compact block streaming is delegated to "
        "lightwalletd. The wallet's responsibilities are key management, note discovery, witness "
        "maintenance, transaction construction, proving, signing, and broadcast coordination.",
    )
    for item in [
        "In scope: Orchard HD wallets, CLI and companion APIs, compact sync, local witness derivation.",
        "Out of scope: Consensus rule changes, t-address support in UI, custodial key hosting as a product goal.",
        "Deferred on roadmap: Secret/Shade/XMR multichain lanes; near-term focus is pure Zcash.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Design Principles", 2)
    for p in [
        "Privacy first — shielded Orchard flows are mandatory, not configurable off.",
        "Zebra-native — JSON-RPC to zebrad; compact blocks via lightwalletd gRPC.",
        "Local witness derivation — spends prove against client-maintained incremental witnesses; anchor from z_gettreestate.",
        "Deterministic indexing — same mnemonic and block range must yield identical note sets.",
        "Minimal trust — servers deliver blocks and relay transactions; they do not hold spending keys.",
    ]:
        add_bullet(doc, p)

    add_heading(doc, "1.4 Development Timeline", 2)
    add_body(
        doc,
        "September 2025: Project inception; HD wallet (BIP39/BIP32), encrypted storage, Orchard "
        "key derivation. October–November 2025: Note scanning, transaction building, Zebra RPC "
        "integration, CLI. December 2025: NU 6.1 support, security hardening (elimination of "
        "panic paths in library code, mutex poisoning guards). Q1 2026: Zeaking compact sync Phase 1, "
        "browser extension Orchard pipeline, dynamic-fee pilot v2.3.0. Q2 2026: NU 6.2 mainnet "
        "compatibility (v2.3.3), Send Select coin selection (v2.3.4), api-server VPS hardening "
        "(v2.3.6.x), mobile companion scaffold. Current release: v2.3.6.5 Teriyaki Hot, June 2026.",
    )
    doc.add_page_break()


def section2_background(doc):
    add_heading(doc, "2. Background: Zcash Shielded Pools", 1)

    add_heading(doc, "2.1 Evolution of Shielded Technology in Zcash", 2)
    add_body(
        doc,
        "Zcash launched with the Sprout shielded pool, which required large parameter ceremonies "
        "and proved expensive to use at scale. Sapling (2018) introduced more efficient proofs and "
        "better performance. Orchard (2021, activated on mainnet at NU5) replaced Sapling as the "
        "recommended shielded pool for new development. Orchard uses the Halo 2 proving system, "
        "which eliminates the need for a per-circuit trusted setup specific to Orchard spends in "
        "the same manner as earlier SNARK schemes required.",
    )
    add_body(
        doc,
        "NozyWallet targets Orchard exclusively for new wallet activity. Legacy Sapling-only flows "
        "are not the product focus. Unified addresses (ZIP-316) bundle Orchard and optionally other "
        "receiver types into a single encoded string (u1 on mainnet); NozyWallet generates and parses "
        "Orchard-capable unified addresses for receive and send.",
    )

    add_heading(doc, "2.2 Why Orchard for a Privacy-First Wallet", 2)
    add_body(
        doc,
        "Orchard provides strong transaction privacy through action-based transactions. Each action "
        "can simultaneously describe a spend and an output, improving symmetry and reducing "
        "metadata leakage compared to older designs that separated spends and outputs more rigidly. "
        "For a wallet that refuses transparent operations, Orchard is the current state of the art "
        "within the Zcash protocol family and aligns with ecosystem direction (NU6.2, librustzcash "
        "maintenance, Zebra support).",
    )

    add_heading(doc, "2.3 Network Upgrades Relevant to NozyWallet", 2)
    rows = [
        ("NU 6.1", "Block 3,146,400 (Nov 2025)", "Protocol version 170140; ZIP-271, ZIP-1016"),
        ("NU 6.2", "Active on mainnet 2026", "Branch ID 0x5437f330; orchard 0.14 dependency set"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Upgrade", "Activation", "Relevance"]):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    doc.add_paragraph()
    add_body(
        doc,
        "NozyWallet v2.3.3+ pins librustzcash crates aligned with NU6.2. Sending with stale branch "
        "IDs produces zebrad rejection (RPC code -25, incorrect consensus branch id). Operators "
        "must run NU6.2-aware Zebrad (5.x) on mainnet.",
    )
    doc.add_page_break()


def section3_orchard(doc):
    add_heading(doc, "3. The Orchard Protocol and NozyWallet", 1)

    add_heading(doc, "3.1 Orchard in Protocol Terms", 2)
    add_body(
        doc,
        "Orchard is Zcash's shielded pool based on the Pasta curves and Halo 2 zero-knowledge "
        "proofs. A transaction may contain an Orchard bundle composed of one or more actions. "
        "Each action references shielded notes: commitments that represent held value, encrypted "
        "such that only holders of appropriate viewing keys can determine recipient and amount.",
    )
    add_body(
        doc,
        "Three concepts recur throughout wallet implementation: notes (owned value), nullifiers "
        "(unique spend markers published when a note is consumed), and commitments (append-only "
        "Merkle tree leaves that anchor the shielded state). NozyWallet must discover notes "
        "belonging to the user, track whether they remain unspent, maintain witnesses for Merkle "
        "inclusion proofs, and construct new actions when the user sends ZEC.",
    )

    add_heading(doc, "3.2 Notes, Commitments, and the Merkle Tree", 2)
    add_body(
        doc,
        "An Orchard note encodes a value (in zatoshi), a recipient address component (diversifier "
        "and transmission key), and randomness (rho, rseed) binding the note to its commitment. "
        "When received, the note is appended to the global Orchard commitment tree. Spending requires "
        "proving inclusion of the note's commitment at a valid anchor height—a root of that tree "
        "frozen at a block boundary.",
    )
    add_body(
        doc,
        "NozyWallet persists discovered notes in notes.json (SerializableOrchardNote). Fields "
        "include value, address bytes, nullifier bytes, block height, txid, spent flag, optional "
        "memo, rho/rseed for canonical nullifier recomputation (v2.3.3+), and orchard_incremental_witness_hex "
        "for local spend proving. Witness blobs are advanced in chain order during scan and again "
        "at spend-build time to the current tip.",
    )

    add_heading(doc, "3.3 Nullifiers and Double-Spend Prevention", 2)
    add_body(
        doc,
        "When an Orchard note is spent, its nullifier is revealed on-chain. Nullifiers are deterministic "
        "functions of the note and the holder's full viewing key; two spends of the same note would "
        "publish the same nullifier, allowing the network to reject duplicates. Wallets must track "
        "nullifiers to exclude spent notes from the spendable set.",
    )
    add_body(
        doc,
        "NozyWallet v2.3.3 (PR #61) corrected a critical bug: compact-block discovery had stored "
        "action nullifiers that did not match canonical note nullifiers, so mark_note_spent failed "
        "and spent notes could be re-selected. The fix derives canonical nullifiers at discovery via "
        "note.nullifier(fvk), persists rho and rseed, marks spends when scanning on-chain action "
        "nullifiers, and records canonical nullifiers on broadcast. Users upgrading should run "
        "nozy sync --to-tip once to repair persisted state.",
    )

    add_heading(doc, "3.4 Actions, Bundles, and Transaction Structure", 2)
    add_body(
        doc,
        "Orchard transactions group logic into actions within a bundle. A typical NozyWallet send "
        "uses build_single_spend: one spend action consuming a selected note, one or two outputs "
        "(recipient and optional change). ZIP-317 fee calculation counts logical actions as "
        "max(spends, outputs) within the Orchard bundle, not the sum—v2.3.3 corrected overpayment "
        "from additive counting.",
    )
    add_body(
        doc,
        "The transaction builder (transaction_builder.rs, orchard_tx.rs) orchestrates: coin selection "
        "(select_single_spend_note chooses the smallest note covering amount plus fee), witness "
        "refresh, proof generation via Halo 2, signing with SpendingKey, serialization to raw "
        "hex, and broadcast through ZebraClient.sendrawtransaction.",
    )

    add_heading(doc, "3.5 Halo 2 Proving in NozyWallet", 2)
    add_body(
        doc,
        "Orchard spends require zero-knowledge proofs that the spender knows valid opening "
        "information for an unspent note, that values balance, and that nullifiers are correctly "
        "derived—without revealing which note or which address. NozyWallet uses the orchard crate "
        "and zcash_proofs for proving parameters. Proving occurs locally on the device running "
        "the wallet (CLI, api-server host, or WASM in extension); it is not outsourced to the node.",
    )
    add_body(
        doc,
        "Proving is computationally heavier than transparent signing. NozyWallet suppresses verbose "
        "scan logging by default (NOZY_VERBOSE_SCAN=1 for detail). Operators should expect "
        "noticeable CPU use during first send after sync on constrained hardware.",
    )

    add_heading(doc, "3.6 Unified Addresses (ZIP-316)", 2)
    add_body(
        doc,
        "Users interact with encoded unified addresses (u1… on mainnet). Internally, NozyWallet "
        "derives Orchard addresses from ZIP-32: SpendingKey::from_zip32_seed(seed, coin_type 133, "
        "account_id). HDWallet.generate_orchard_address produces diversifier indices for fresh "
        "receive addresses. Validation accepts u1 strings up to 256 characters (v2.3.5 fixed "
        "api-server rejection of ~106-character valid mainnet UAs).",
    )

    add_heading(doc, "3.7 Note Discovery and Decryption", 2)
    add_body(
        doc,
        "Discovery scans blocks from wallet birthday (or last_scan_height) to node tip. For each "
        "Orchard action in compact or full blocks, NozyWallet attempts trial decryption with "
        "IncomingViewingKey (external and internal scopes). Successful decryption yields note "
        "plaintext, value, memo field, and metadata for persistence.",
    )
    add_body(
        doc,
        "Implementation lives primarily in src/notes.rs (NoteScanner) and integrates with "
        "NoteIndex for deterministic ordering. Compact sync via zeaking::lwd replays cached "
        "compact blocks; RPC sync fetches via zebrad getblock. Both paths must converge on "
        "identical note sets for a given range—verified by deterministic_scanning_tests.rs.",
    )

    add_heading(doc, "3.8 Witness Derivation and the Send Pipeline", 2)
    add_body(
        doc,
        "A Merkle witness proves a note commitment exists in the Orchard tree at anchor height. "
        "Early NozyWallet development hit a roadblock: depending on Zebrad JSON-RPC for witness "
        "paths (position/authpath style data) was unreliable for proving. The current architecture "
        "documented in ZEBRAD_SHIELDED_SEND_LIMIT.md requires:",
    )
    for step in [
        "Persist incremental witnesses per note during scan (OrchardWitnessTracker).",
        "Advance witnesses block-by-block in chain order to spend height.",
        "Fetch anchor from z_gettreestate at spend height; verify witness root matches anchor.",
        "Supply witness to Orchard proof builder; sign and broadcast.",
    ]:
        add_numbered(doc, step)
    add_body(
        doc,
        "load_spendable_notes_from_wallet (notes.rs) loads cached notes for send without a 50k-block "
        "rescan when witnesses are present—v2.3.6.x fix for api-server send latency. Fallback scan "
        "uses incremental bounds when cache is empty.",
    )

    add_heading(doc, "3.9 Memos and ZIP-271", 2)
    add_body(
        doc,
        "Orchard outputs may carry encrypted memo fields (512-byte chunks affect ZIP-317 logical "
        "action counts). NozyWallet supports memos on send where transaction size limits allow. "
        "Business roadmap (Phase 4) includes optional invoice identifiers in memos for vendor flows.",
    )

    add_heading(doc, "3.10 Orchard vs Transparent: NozyWallet Policy", 2)
    add_body(
        doc,
        "The wallet never presents t1 addresses for receive in normal flows and rejects them on "
        "send. This policy is stricter than the Zcash protocol, which remains dual-pool. Users "
        "who must interact with transparent-only services need a different wallet. NozyWallet "
        "optimizes for users who want Orchard-only operational semantics.",
    )
    doc.add_page_break()


def section4_viewing_keys(doc):
    add_heading(doc, "4. Viewing Keys in Zcash and NozyWallet", 1)

    add_heading(doc, "4.1 Purpose of Viewing Keys", 2)
    add_body(
        doc,
        "Shielded transactions hide data from the public ledger, but the wallet holder—and optionally "
        "third parties they authorize—must still decrypt incoming and outgoing activity. Zcash "
        "separates spending authority from viewing authority. Spending keys can sign transactions "
        "and derive nullifiers; viewing keys can decrypt note ciphertexts and observe balances but "
        "cannot spend funds. NozyWallet uses viewing material internally on every sync and exposes "
        "controlled export paths for hardware pairing and (planned) accounting disclosure.",
    )

    add_heading(doc, "4.2 Key Hierarchy in NozyWallet", 2)
    add_body(
        doc,
        "From the user's BIP39 mnemonic, NozyWallet derives a ZIP-32 unified spending key (USK) "
        "per account index (coin type 133). From the USK:",
    )
    rows = [
        ("Unified Spending Key (USK)", "Full spend authority; never exported in disclosure flows"),
        ("Unified Full Viewing Key (UFVK)", "ZIP-316 encoded; Orchard + optional other components"),
        ("Orchard Full Viewing Key (FVK)", "Orchard-specific viewing; nullifier derivation"),
        ("Incoming Viewing Key (IVK)", "External/internal scopes; trial decryption when scanning"),
        ("SpendingKey", "Orchard spend authority for proof and signature"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Key type"
    table.rows[0].cells[1].text = "Role in NozyWallet"
    for ri, row in enumerate(rows, 1):
        table.rows[ri].cells[0].text = row[0]
        table.rows[ri].cells[1].text = row[1]
    doc.add_paragraph()

    add_heading(doc, "4.3 Internal Use During Scan", 2)
    add_body(
        doc,
        "On every sync, NozyWallet derives FullViewingKey and IncomingViewingKey values from the "
        "unlocked HD wallet. NoteScanner iterates compact or full Orchard actions and attempts "
        "decryption (try_compact_note_decryption for compact blocks). Matching actions become "
        "persisted notes. Canonical nullifiers are computed with note.nullifier(fvk) at discovery.",
    )
    add_body(
        doc,
        "This process never uploads viewing keys to Zebrad or lightwalletd. The node supplies "
        "public chain data; decryption is entirely client-side. A hostile RPC operator learns only "
        "that a client requested blocks—not which notes belonged to the wallet.",
    )

    add_heading(doc, "4.4 UFVK Export: Keystone Hardware Wallet", 2)
    add_body(
        doc,
        "NozyWallet implements Keystone integration (src/keystone.rs, api-server /api/keystone/*, "
        "mobile Keystone screen). export_ufvk_from_wallet derives the unified full viewing key "
        "from the mnemonic (currently account 0) and encodes it for device pairing. The hardware "
        "device holds spend approval; the host prepares unsigned PCZT (Partially Created Zcash "
        "Transaction), redacting witness paths and derivation metadata before UR/QR transfer.",
    )
    add_body(
        doc,
        "PCZT redaction (keystone.rs) clears spend witnesses, zip32 derivation fields, dummy "
        "spending keys, and proprietary metadata from Orchard actions before the user scans with "
        "Keystone. After device signature, complete_send broadcasts via RPC. This split keeps "
        "spending keys off the networked host while allowing the host to perform scanning and proof "
        "generation that requires witness data.",
    )

    add_heading(doc, "4.5 Viewing Keys vs Spending Keys: Security Boundary", 2)
    add_body(
        doc,
        "Users and operators must treat UFVK export as sensitive. Anyone with a UFVK for an account "
        "can view all past and future shielded activity visible to that account's IVKs—they cannot "
        "steal funds, but privacy against that party is lost retroactively for the covered scope. "
        "NozyWallet policy (AGENTS.md): never log seeds, mnemonics, or raw spending keys; UFVK "
        "export requires explicit user action.",
    )

    add_heading(doc, "4.6 Planned Selective Disclosure (Business Phase 5)", 2)
    add_body(
        doc,
        "The pure-Zcash business roadmap (GitHub issue #85) separates payment identity (Zcash Names, "
        "e.g. hotwing.zcash) from accounting visibility (viewing keys). Planned Phase 5 features:",
    )
    for item in [
        "WalletProfile::Business uses Orchard ZIP-32 account index 1; Personal uses account 0.",
        "CSV ledger export from merged transaction history (no third-party key required for owner).",
        "nozy disclosure export / API: export Business-account UFVK only, never USK or mnemonic.",
        "Optional scoped IVK + height range for tighter disclosure (design TBD).",
        "Disclosure grant log: append-only record of grant id, scope, creation, expiry, access events.",
        "Settings UX: Share with accountant with explicit warnings about retroactive visibility.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.7 Viewing Keys and ZNS: Orthogonal Layers", 2)
    add_body(
        doc,
        "Zcash Names (ZNS) resolves human-readable names to unified addresses via an indexer. "
        "Customers paying hotwing.zcash do not receive or need viewing keys. ZNS claim/update "
        "(Phase 6) uses a separate Ed25519 identity key for on-chain name registration memos—distinct "
        "from Orchard UFVK/IVK material. Confusing name identity keys with viewing keys would mislead "
        "users about who can see transaction history.",
    )

    add_heading(doc, "4.8 Comparison: What Each Party Holds", 2)
    rows = [
        ("Customer", "Nothing; pays to resolved UA or ZIP-321 URI"),
        ("Vendor (owner)", "Mnemonic + password; full spend and view on Business account"),
        ("Accountant", "Exported Business UFVK (read-only); no spend"),
        ("Public indexer", "Name → address mapping only"),
        ("RPC operator", "Block data; no wallet keys unless user self-hosts compromised host"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Party"
    table.rows[0].cells[1].text = "Capability"
    for ri, row in enumerate(rows, 1):
        table.rows[ri].cells[0].text = row[0]
        table.rows[ri].cells[1].text = row[1]
    doc.add_paragraph()

    add_heading(doc, "4.9 Implementation Status Summary", 2)
    rows = [
        ("IVK/FVK scan decryption", "Shipped — core sync path"),
        ("UFVK export for Keystone", "Shipped — account 0"),
        ("UFVK export for Business account", "Planned — Phase 5"),
        ("Disclosure grant audit log", "Planned — Phase 5"),
        ("In-wallet ZNS identity key", "Planned — Phase 6"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Capability"
    table.rows[0].cells[1].text = "Status"
    for ri, row in enumerate(rows, 1):
        table.rows[ri].cells[0].text = row[0]
        table.rows[ri].cells[1].text = row[1]
    doc.add_paragraph()
    doc.add_page_break()


def section5_architecture(doc):
    add_heading(doc, "5. System Architecture", 1)

    add_heading(doc, "5.1 Component Overview", 2)
    stack_rows = [
        ("nozy (CLI + lib)", "src/", "HD wallet, notes, orchard_tx, fee_policy, ZebraClient"),
        ("zeaking", "zeaking/", "Compact sync, SQLite cache, LWD gRPC client"),
        ("nozywallet-api", "api-server/", "REST companion, sync/send/history endpoints"),
        ("desktop-client", "desktop-client/", "Tauri 2 + React GUI"),
        ("browser-extension", "browser-extension/", "MV3 service worker + WASM core"),
        ("nozy-mobile", "nozy-mobile/", "Expo app → api-server"),
        ("zeaking-ffi", "zeaking-ffi/", "UniFFI bindings for mobile native sync"),
    ]
    table = doc.add_table(rows=1 + len(stack_rows), cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Component", "Location", "Function"]):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(stack_rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    doc.add_paragraph()

    add_heading(doc, "5.2 Data Flow: Receive", 2)
    for step in [
        "User shares unified address (u1) or future ZNS name (hotwing.zcash).",
        "Sender constructs Orchard transaction externally; note commitment appended on-chain.",
        "NozyWallet sync scans blocks; IVK trial decryption finds matching action.",
        "Note persisted to notes.json with witness advancement; balance updated.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "5.3 Data Flow: Send", 2)
    for step in [
        "User specifies recipient (u1 or resolved ZNS name) and amount.",
        "Wallet loads spendable notes from cache; select_single_spend_note picks covering note.",
        "Witness advanced to tip; anchor verified via z_gettreestate.",
        "Halo 2 proof generated; transaction signed and broadcast.",
        "Spent note marked in notes.json; history records outgoing tx.",
    ]:
        add_numbered(doc, step)

    add_heading(doc, "5.4 Sync Paths", 2)
    add_body(
        doc,
        "RPC sync (nozy sync, POST /api/sync): walks blocks via zebrad, merges into notes.json, "
        "updates last_scan_height. Compact sync (zeaking::lwd): streams GetBlockRange from "
        "lightwalletd into lwd_compact.sqlite; supports resume; surfaces structured errors "
        "(LWD_GRPC, LWD_STORAGE). nozy sync --to-tip scans through chain tip—required after "
        "receives on mainnet because plain sync advances ~1000 blocks per invocation.",
    )

    add_heading(doc, "5.5 Operator Stack", 2)
    add_body(
        doc,
        "Recommended deployment: zebrad (JSON-RPC :8232) + lightwalletd (gRPC :9067) + nozy CLI "
        "or nozywallet-api. Mobile vendors may use hosted api-server on VPS "
        "(nozy-mobile/VPS-DEPLOY.md) pointing at remote Zebrad. trusted_zebra_urls (v2.3.6.3) "
        "allows direct RPC to operator VPS when privacy-network policy would otherwise require Tor.",
    )
    doc.add_page_break()


def section6_features(doc):
    add_heading(doc, "6. Implemented Features and Release History", 1)

    add_heading(doc, "6.1 Core CLI Capabilities", 2)
    for f in [
        "nozy create / restore — BIP39 wallet with optional Argon2 password",
        "nozy receive — display unified Orchard address",
        "nozy send [--priority] — shielded send with ZIP-317 fee",
        "nozy sync [--to-tip] — incremental and full tip sync",
        "nozy balance / status / history — wallet state and diagnostics",
        "nozy lwd sync-to-tip / prune — compact cache maintenance",
        "nozy nu61 — NU 6.1 status helper",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "6.2 Zeaking and Compact Sync", 2)
    add_body(
        doc,
        "Phase 1 (May 2026) shipped resume-capable compact sync, structured gRPC errors, "
        "api-server /api/lwd/sync/compact, desktop and FFI bindings, extension companion resume "
        "flag. Planned: witness advancement from compact bytes to reduce getblock dependency.",
    )

    add_heading(doc, "6.3 Dynamic-Fee Pilot", 2)
    add_body(
        doc,
        "Aligned with Shielded Labs pilot shape: ZIP-317 conventional fees, opt-in 4× priority "
        "multiplier, nExpiryHeight = chain_tip + 1 + 5 blocks (~6 min at 75 s/block). Speed-up "
        "rebuilds expired txs at priority fee. Not mempool-driven congestion pricing—client-side "
        "policy only.",
    )

    add_heading(doc, "6.3.1 Proving Latency vs Pilot Expiry (BUG-2026-011)", 3)
    add_body(
        doc,
        "Two clocks must not be conflated: (1) build clock—witness fetch, Orchard Halo2 prove, "
        "sign, broadcast; (2) mempool expiry clock—blocks after successful broadcast until "
        "nExpiryHeight, when the wallet marks a tx Expired and offers speed-up. VPS testing "
        "(Gilmore, June 2026) showed proving can span multiple blocks; encoding expiry from an "
        "early chain tip caused Zebrad -25 rejections before mempool admission even though the tx "
        "was fully signed.",
    )
    add_body(
        doc,
        "Fix without lengthening expiry: refresh chain tip immediately before encoding "
        "nExpiryHeight; auto-rebuild (up to 3 attempts) when proving outruns the window; retry "
        "broadcast on expiry consensus errors. PILOT_EXPIRY_DELTA_BLOCKS remains 5.",
    )
    add_body(
        doc,
        "Why 5 blocks beats 15: a 15-block delta (~19 min) would defer expire-and-replace feedback "
        "for all users—users wait too long to learn a send failed or to speed up. Proving latency "
        "is an implementation/runtime problem (refresh + rebuild); expiry delta is a product knob "
        "for confirmation feedback. See docs/reference/PILOT_EXPIRY_PROVING_LATENCY.md.",
    )

    add_heading(doc, "6.3.2 Send Readiness and Mainnet Evidence (June 2026)", 3)
    add_body(
        doc,
        "Operator mainnet runs (WSL Zebrad, Windows CLI) recorded: sync of 5132 blocks in ~32 s; "
        "witness lag 1 block after sync; shielded sends ~198–206 s end-to-end with broadcast success "
        "and no -25 expiry; stale wallet (~5000 blocks behind) rejected in ~0.09 s with sync-to-tip "
        "message instead of 7+ minute witness catch-up at send time. Mechanisms: "
        "MAX_SEND_WITNESS_LAG_BLOCKS=50, parallel getblock catch-up (10/batch), proving warm-up "
        "~2 s cold, NoteIndex v2 mark-spent after broadcast. TXIDs: 5a03fbd1…, 902cf006…. "
        "Full tables: docs/reference/MAINNET_SEND_READINESS_EVIDENCE.md.",
    )

    add_heading(doc, "6.4 API Server Highlights", 2)
    for a in [
        "Wallet lifecycle: create, unlock, lock, status",
        "POST /api/sync — unified wallet_sync orchestrator (v2.3.6.2)",
        "POST /api/transaction/send — cache-first note loading",
        "GET /api/transaction/history — merged sent + received (master, post-v2.3.6.5)",
        "POST /api/transaction/speed-up — expired tx rebuild",
        "Keystone: export-ufvk, prepare-send, complete-send",
    ]:
        add_bullet(doc, a)

    add_heading(doc, "6.5 Release Milestones", 2)
    milestones = [
        ("v2.3.0", "Priority Lane — ZIP-317 fees, --priority, 2-block expiry"),
        ("v2.3.1", "sync --to-tip, lwd maintenance commands"),
        ("v2.3.3", "NU6.2, 5-block expiry, fee action fix, nullifier fix (#58–#61)"),
        ("v2.3.4", "Send Select — multi-note coin selection"),
        ("v2.3.6", "CLI-only GitHub releases; workspace hardening"),
        ("v2.3.6.5", "Empty-cache rescan fix; VPS connect policy"),
    ]
    table = doc.add_table(rows=1 + len(milestones), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Version"
    table.rows[0].cells[1].text = "Highlight"
    for ri, row in enumerate(milestones, 1):
        table.rows[ri].cells[0].text = row[0]
        table.rows[ri].cells[1].text = row[1]
    doc.add_paragraph()

    add_heading(doc, "6.6 Surfaces in Development", 2)
    add_body(
        doc,
        "Desktop: History CSV, address book, onboarding, fee toggle, speed-up. Extension: WASM "
        "Orchard v5 build, witness RPC inputs, broadcast, dApp provider. Mobile: Expo screens "
        "for wallet lifecycle, Keystone, hosted API URL configuration.",
    )
    doc.add_page_break()


def section7_security(doc):
    add_heading(doc, "7. Security, Privacy, and Deterministic Indexing", 1)

    add_heading(doc, "7.1 Cryptographic Privacy Properties", 2)
    add_body(
        doc,
        "Orchard hides sender identity, receiver identity, and amount from public ledger observers "
        "who lack viewing keys. NozyWallet adds product-layer enforcement by blocking transparent "
        "addresses. Users still must consider network metadata (RPC timing, IP address), device "
        "compromise, and backup security.",
    )

    add_heading(doc, "7.2 Storage and Secrets", 2)
    for k in [
        "Encrypted wallet files with Argon2 password hashing",
        "zeroize for sensitive buffers in library paths",
        "No telemetry of addresses, amounts, or seeds (AGENTS.md)",
        "API key auth for public VPS deployments (operator responsibility)",
    ]:
        add_bullet(doc, k)

    add_heading(doc, "7.3 Deterministic Indexing", 2)
    add_body(
        doc,
        "ShieldOrder (Zcash community review, Dec 2025) emphasized reproducible indexing. NozyWallet "
        "implements NoteIndex ordering, ascending block scan, nullifier deduplication, and ignored "
        "integration tests over fixed block ranges. Same mnemonic + range → identical note set.",
    )

    add_heading(doc, "7.4 Mainnet Incident Response", 2)
    add_body(
        doc,
        "Field testing exposed branch ID mismatch, aggressive expiry, incorrect fee action counting, "
        "nullifier bugs, and multi-note IncorrectFee—all addressed in v2.3.3–v2.3.4. Community "
        "VPS testing (Gilmore, 2026) exposed send rescan, empty history, and pre-broadcast expiry "
        "rejections on slow proving stacks (BUG-2026-011)—fixed via late tip refresh and rebuild "
        "while keeping 5-block pilot expiry for speed-up UX. This feedback loop is documented in "
        "CHANGELOG.md, docs/reference/PILOT_EXPIRY_PROVING_LATENCY.md, and forum updates.",
    )
    doc.add_page_break()


def section8_roadmap(doc):
    add_heading(doc, "8. Product Roadmap: Pure Zcash Stack", 1)

    add_heading(doc, "8.1 Strategic Pivot", 2)
    add_body(
        doc,
        "Near-term engineering pauses Secret/Shade/XMR multichain lanes and concentrates on Zebrad + "
        "lightwalletd + Orchard + Zeaking + mobile + ZNS. Multichain code may remain for contributors "
        "but is not the public product story.",
    )

    add_heading(doc, "8.2 Phased Delivery", 2)
    phases = [
        ("Phase 0", "Maintainer alignment — GitHub issue #85 (filed)"),
        ("Phase 1", "Personal / Business profile; Orchard account 0 vs 1"),
        ("Phase 2", "ZNS resolve — send to name.zcash"),
        ("Phase 3", "Mobile Sell mode — QR POS (hotwing.zcash demo)"),
        ("Phase 4", "ZIP-321 scan-to-pay URIs"),
        ("Phase 5", "Business CSV + UFVK accountant disclosure"),
        ("Phase 6", "In-wallet ZNS claim/update"),
    ]
    table = doc.add_table(rows=1 + len(phases), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Phase"
    table.rows[0].cells[1].text = "Deliverable"
    for ri, row in enumerate(phases, 1):
        table.rows[ri].cells[0].text = row[0]
        table.rows[ri].cells[1].text = row[1]
    doc.add_paragraph()

    add_heading(doc, "8.3 Remaining Technical Debt", 2)
    for r in [
        "Zeaking witness from LWD compact bytes",
        "Extension dynamic-fee checklist completion",
        "Third-party security audit (2026 target)",
        "Business UFVK export security review before Phase 5 ship",
    ]:
        add_bullet(doc, r)
    doc.add_page_break()


def section9_conclusion(doc):
    add_heading(doc, "9. Conclusion", 1)
    add_body(
        doc,
        "NozyWallet implements a shielded-first Zcash wallet on modern infrastructure: Orchard "
        "cryptography via librustzcash, Zebrad for node RPC, lightwalletd for compact blocks, and "
        "local witness derivation for reliable sends. This report detailed Orchard protocol concepts "
        "— notes, nullifiers, witnesses, actions, Halo 2 proofs — and how viewing keys enable "
        "scanning, hardware signing, and planned business disclosure without conflating them with "
        "ZNS identity or spending authority.",
    )
    add_body(
        doc,
        "From September 2025 through v2.3.6.5, the project moved from foundation to mainnet-validated "
        "operation with community-tested operator paths. The roadmap ahead is intentionally narrow: "
        "pure Zcash usability for individuals and small vendors, with human-readable names and "
        "optional read-only accounting exports. Continued contribution, VPS operator feedback, and "
        "formal review before disclosure features ship remain essential.",
    )

    add_heading(doc, "References", 1)
    refs = [
        "LEONINE-DAO/Nozy-wallet — github.com/LEONINE-DAO/Nozy-wallet",
        "Zebra — github.com/ZcashFoundation/zebra",
        "lightwalletd — github.com/zcash/lightwalletd",
        "Orchard crate — github.com/zcash/orchard",
        "ZIP-224/225 Orchard — zips.z.cash",
        "ZIP-316 Unified addresses — zips.z.cash",
        "ZIP-317 Conventional fees — zips.z.cash",
        "Zcash Names — zcashnames.com/docs",
        "ZEBRAD_SHIELDED_SEND_LIMIT.md — NozyWallet repo",
        "Issue #85 — Business / ZNS tracking",
    ]
    for r in refs:
        add_bullet(doc, r)

    add_heading(doc, "Appendix A: Glossary", 1)
    glossary = [
        ("Action", "Orchard bundle unit combining spend and output semantics"),
        ("Anchor", "Orchard commitment tree root at a block height"),
        ("Compact block", "Lightwalletd compressed block format for efficient sync"),
        ("FVK", "Full viewing key — Orchard viewing + nullifier derivation"),
        ("IVK", "Incoming viewing key — decrypts received note ciphertexts"),
        ("Nullifier", "Unique value revealed when a note is spent"),
        ("PCZT", "Partially Created Zcash Transaction — hardware signing format"),
        ("UFVK", "Unified full viewing key — ZIP-316 encoded viewing key"),
        ("Witness", "Merkle inclusion proof for a note commitment"),
        ("ZNS", "Zcash Name System — on-chain name to address mapping"),
    ]
    table = doc.add_table(rows=1 + len(glossary), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    for ri, row in enumerate(glossary, 1):
        table.rows[ri].cells[0].text = row[0]
        table.rows[ri].cells[1].text = row[1]
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.italic = True
    run.font.size = Pt(10)


def main():
    doc = Document()
    set_margins(doc.sections[0])

    cover_page(doc)
    table_of_contents(doc)
    abstract(doc)
    section1_intro(doc)
    section2_background(doc)
    section3_orchard(doc)
    section4_viewing_keys(doc)
    section5_architecture(doc)
    section6_features(doc)
    section7_security(doc)
    section8_roadmap(doc)
    section9_conclusion(doc)

    doc.save(OUTPUT)
    para_count = len(doc.paragraphs)
    print(f"Saved: {OUTPUT}")
    print(f"Paragraphs: {para_count} (approx. 18–24 pages in Word at 11pt)")


if __name__ == "__main__":
    main()
