#!/usr/bin/env python3
"""Generate NozyWallet white paper (Word + Markdown mirror).

Content covers Orchard→Ironwood, ZIP-317 / ZIP 318 math, Nym network privacy,
and release history through the v2.4.1 Ironwood era.

  python scripts/generate-nozy-whitepaper.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = REPO_ROOT / "docs" / "NozyWallet_Whitepaper.docx"
OUTPUT_MD = REPO_ROOT / "docs" / "reference" / "NozyWallet_Whitepaper.md"
VERSION = "2.4.1 (Ironwood era)"


def set_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(3)


def add_adr(doc, title, context, decision, consequences):
    add_heading(doc, title, 3)
    add_body(doc, f"Context: {context}")
    add_body(doc, f"Decision: {decision}")
    add_body(doc, f"Consequences: {consequences}")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, 1):
        for ci, cell in enumerate(row):
            table.rows[ri].cells[ci].text = str(cell)
    doc.add_paragraph()


def add_formula(doc, text):
    """Monospace-ish body line for equations (Word-friendly)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Consolas"
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)


def cover_page(doc):
    logo_path = REPO_ROOT / "docs" / "reference" / "NozyWallet_Whitepaper_logo.png"
    if logo_path.is_file():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Inches(2.6))
        p_logo.paragraph_format.space_after = Pt(12)
    else:
        for _ in range(2):
            doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NozyWallet")
    r.bold = True
    r.font.size = Pt(36)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        "White Paper\n"
        "Architecture, Shielded Pools, Network Privacy, Zcash Names, and Mainnet Operation"
    )
    r2.font.size = Pt(14)
    r2.font.name = "Calibri"
    r2.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "LEONINE DAO",
        f"Version {VERSION}",
        date.today().strftime("%B %Y"),
        "",
        "github.com/LEONINE-DAO/Nozy-wallet",
        "nozywallet.com",
    ]:
        run = meta.add_run(line + "\n")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    doc.add_page_break()


def section_executive_summary(doc):
    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "NozyWallet is a self-custodial, shielded-first Zcash wallet and companion stack—not a "
        "consensus node. Operators run Zebrad for JSON-RPC (broadcast, tip, treestate) and "
        "lightwalletd for compact-block sync. The wallet derives Merkle witnesses locally, "
        "computes ZIP-317 fees on-device, and builds ZIP-225 / Ironwood transactions with Halo 2 "
        "proofs on the user’s machine.",
    )
    add_body(
        doc,
        "What ships today (v2.4.x): mainnet-validated CLI (v2.4.1 / v2.4.1.1), localhost "
        "api-server, desktop 1.0.0-beta.2, browser extension beta, and mobile companion "
        "surfaces—sharing the nozy Rust core and zeaking compact sync. Nozy participates in the "
        "Shielded Labs dynamic-fee pilot and ships Ironwood (NU6.3) migration tooling aligned "
        "with draft ZIP 318.",
    )
    add_body(
        doc,
        "Why Ironwood: a 2026 Orchard circuit soundness issue motivated sealing Orchard and "
        "opening Ironwood with a corrected circuit and turnstile so circulating supply can be "
        "reasoned about again. After activation height 3,428,143 (2026-07-28), ordinary sends "
        "use Ironwood; Orchard value migrates via Plan → Split → Migrate → Broadcast.",
    )
    add_body(
        doc,
        "Network metadata: shielded proofs hide note contents; they do not hide that an IP "
        "submitted a transaction. Nozy’s Priority-1 path is a hybrid—baseline hygiene plus "
        "optional Nym mixnet broadcast / dVPN sync—without claiming full “Nym integration” "
        "until live remote submit evidence is complete.",
    )
    add_body(
        doc,
        "This paper states architecture decisions, mathematical policy (fees, expiry, ZIP 318 "
        "buckets and denominations), Ironwood and Nym work, release history, field evidence, "
        "and lessons for operators, Shielded Labs reviewers, grant readers, and contributors.",
    )
    doc.add_page_break()


def section_architecture_decisions(doc):
    add_heading(doc, "1. Architecture Decisions", 1)
    add_body(
        doc,
        "Each decision follows an ADR pattern: context, decision, and consequences.",
    )
    add_adr(
        doc,
        "1.1 Zebrad-only stack (no zcashd)",
        "Zcash infrastructure consolidates on Zebra; pilots target Zebrad + lightwalletd.",
        "Zebrad JSON-RPC for broadcast, blocks, and treestate; lightwalletd gRPC via "
        "zeaking::lwd. No zcashd in this repository.",
        "No estimatefee—fees are client-side. No spend-ready witness RPC—witnesses are local.",
    )
    add_adr(
        doc,
        "1.2 One Rust core, multiple surfaces",
        "Fee, expiry, witness catch-up, and broadcast retry must match across surfaces.",
        "Centralize logic in nozy; thin CLI, api-server, desktop, extension, and mobile.",
        "Extension WASM is workspace-excluded. Surface drift caused BUG-2026-001–011.",
    )
    add_adr(
        doc,
        "1.3 Local witness derivation",
        "Shielded spends need Merkle paths and anchors consistent with chain state.",
        "Persist incremental witnesses; catch up via Zebra blocks; verify with z_gettreestate.",
        "Reject sends if witness lag L > 50 blocks; require sync-to-tip first.",
    )
    add_adr(
        doc,
        "1.4 Client-side ZIP-317 fees",
        "Pilot needs deterministic fees; Zebrad lacks fee estimation RPCs.",
        "fee_policy.rs: ZIP-317 conventional fee; priority multiplier m = 4 on all sends.",
        "Logical actions use max(spends, outputs), not sum (v2.3.2 fix).",
    )
    add_adr(
        doc,
        "1.5 Five-block pilot expiry",
        "Short mempool life enables Expired → priority rebuild.",
        "Δ_exp = 5; h_expiry = h_tip + 1 + Δ_exp. Rebuild ≤3 attempts—not Δ_exp = 15.",
        "Fifteen-block expiry reverted (a72bc6e8): slowed speed-up UX.",
    )
    add_adr(
        doc,
        "1.6 Note index v2",
        "Fast load, merged history, mark-spent after broadcast.",
        "Versioned NoteIndex with nullifier/height maps; atomic rename writes.",
        "Legacy array parsers broke mark-spent until June 2026 fixes.",
    )
    add_adr(
        doc,
        "1.7 Ironwood-first after NU6.3",
        "Orchard sealed after Ironwood activation.",
        "Route new sends to Ironwood; Orchard→Ironwood turnstile (ZIP 318, {1,2,5}×10^k).",
        "Migration is a privacy operation (amounts visible on turnstile).",
    )
    add_adr(
        doc,
        "1.8 Network metadata privacy (Nym hybrid + hygiene)",
        "Clearnet submit links IP to mempool arrival (and turnstile amount).",
        "Baseline hygiene always; Priority-1 includes NymMixnetBroadcast; local Zebrad needs no mixnet.",
        "Do not market “Nym integrated” from attestation alone.",
    )
    add_heading(doc, "1.9 System stack", 3)
    add_formula(
        doc,
        "CLI · api-server · desktop · extension · mobile → nozy + zeaking::lwd + ironwood/* "
        "+ nym_mixnet_broadcast (opt-in) → lightwalletd :9067 · Zebrad :8232",
    )
    doc.add_page_break()


def section_math(doc):
    add_heading(doc, "2. Mathematical Model", 1)
    add_body(
        doc,
        "Constants match src/fee_policy.rs and src/ironwood/migration.rs.",
    )

    add_heading(doc, "2.1 ZIP-317 conventional fee", 2)
    add_body(
        doc,
        "Let f_m = 5000 zatoshis be the marginal fee per logical action, and g = 2 the grace "
        "action count. Logical actions for a shielded send shape:",
    )
    add_formula(doc, "a = max(g, max(n_spends, n_outputs) + a_memo)")
    add_body(
        doc,
        "where a_memo counts memo chunks beyond the two free 512-byte chunks. Conventional and "
        "Nozy fees:",
    )
    add_formula(doc, "F_conv = f_m · a")
    add_formula(doc, "F_Nozy = m · F_conv,  m = 4")
    add_body(
        doc,
        "Floor when a = g = 2: F_conv = 10,000 zat, F_Nozy = 40,000 zat.",
    )

    add_heading(doc, "2.2 Pilot expiry height", 2)
    add_formula(doc, "h_build = h_tip + 1")
    add_formula(doc, "h_expiry = h_build + Δ_exp = h_tip + 6   (Δ_exp = 5)")
    add_body(
        doc,
        "Expired when h_tip > h_expiry. At mean block time t_b ≈ 75 s: "
        "T_exp ≈ Δ_exp · t_b ≈ 375 s ≈ 6.25 min. Fifteen-block policy (~19 min) was rejected.",
    )

    add_heading(doc, "2.3 Witness lag guard", 2)
    add_formula(doc, "L = h_tip − h_w")
    add_formula(doc, "allow send  ⟺  L ≤ L_max = 50")
    add_body(
        doc,
        "Latency model: T_send ≈ T_witness + T_setup + T_Halo2 + T_sign + T_broadcast. "
        "Synced wallets (L ≤ 50) observed T_send ~ 200 s on WSL Zebrad + Windows CLI (June 2026).",
    )

    add_heading(doc, "2.4 ZIP 318 anchor buckets", 2)
    add_body(doc, "Draft ZIP 318 bucket width B = 256 blocks:")
    add_formula(doc, "b(h) = B · floor(h / B),   b_next(h) = b(h) + B")
    add_body(
        doc,
        "T_bucket ≈ 256 · 75 s ≈ 5.33 h. Transfer expiry aligns with B. Default K_max = 4 "
        "same-denomination sends per bucket (until ZIP finalizes).",
    )

    add_heading(doc, "2.5 Migration denominations {1,2,5}×10^k", 2)
    add_body(
        doc,
        "Active Shielded Labs / Appendix A ladder: amounts v ∈ {1,2,5} × 10^k (zatoshis; "
        "smallest practical bucket 0.001 ZEC = 10^5 zat). Residuals r < r_min = 10^5 zat are "
        "abandoned (ZOOKO_RESIDUAL_ABANDON_ZAT). ZIP 318 power-of-ten remains compatibility-only.",
    )

    add_heading(doc, "2.6 Baseline network hygiene", 2)
    add_table(
        doc,
        ["Parameter", "Value", "Role"],
        [
            ("Checkpoint spacing", "256 blocks", "Aligns with ZIP 318 buckets"),
            ("Max overlap rewind", "128 blocks", "Bounds start-height choice"),
            ("Broadcast delay", "U[30, 300] s", "Breaks sync-then-submit timing"),
            ("Tip-sync guard", "120 s", "Refuse migrate-broadcast if tip sync too recent"),
        ],
    )
    doc.add_page_break()


def section_ironwood(doc):
    add_heading(doc, "3. Ironwood (NU6.3) and Migration", 1)
    add_heading(doc, "3.1 Motivation", 2)
    add_body(
        doc,
        "Orchard hides amounts. A circuit soundness bug (Taylor Hornby / Shielded Labs) created "
        "a counterfeiting risk inside the shielded pool. Patching stops new abuse; it does not "
        "by itself restore verifiable circulating supply. Ironwood seals ordinary Orchard "
        "activity, opens a corrected-circuit pool, and moves value through a turnstile.",
    )
    add_heading(doc, "3.2 Activation", 2)
    add_table(
        doc,
        ["Network", "Height", "Calendar target"],
        [("Mainnet NU6.3", "3,428,143", "2026-07-28")],
    )
    add_heading(doc, "3.3 Wallet flow", 2)
    for step in [
        "Sync to tip (witnesses + notes).",
        "Plan ZIP 318 schedule (splits, denominations, buckets).",
        "Split Orchard→Orchard into schedule-sized notes if needed.",
        "Migrate (prebuild locked turnstile transactions).",
        "Broadcast inside the scheduled window with network-privacy mode selected.",
    ]:
        add_bullet(doc, step)
    add_body(
        doc,
        "Surfaces: CLI Ironwood commands, desktop Ironwood UI, mobile screens, landing "
        "/ironwood dashboard, api-server handlers.",
    )
    add_heading(doc, "3.4 Privacy expectations", 2)
    add_body(
        doc,
        "Turnstile crossings reveal amounts on-chain by design. Clearnet broadcast can link "
        "amount to network identity. Nozy treats migration as an explicit privacy warning plus "
        "safer egress—not a silent balance upgrade.",
    )
    doc.add_page_break()


def section_nym(doc):
    add_heading(doc, "4. Network Privacy and Nym", 1)
    add_body(
        doc,
        "Shielded cryptography and mixnets solve different problems. Orchard / Ironwood hide "
        "note contents. Nym (and Tor) hide who talked to which endpoint. Ironwood migration "
        "makes the gap acute: turnstile amounts are public by design, so an observer who also "
        "sees IP ↔ submit time can join network identity to a pool-crossing event.",
    )
    add_body(
        doc,
        "Nozy’s network-privacy work follows three priorities after ZIP 318: (1) protect the "
        "broadcasting IP, (2) share cover traffic across cohorts, (3) amount + timing selection. "
        "Nym engineering sits primarily in Priority 1, plus hygiene transport cannot invent.",
    )

    add_heading(doc, "4.1 Threat model", 2)
    add_body(
        doc,
        "Assumed attacker: ISP/VPS egress observation; remote LWD/submit logs of IP + timestamp + "
        "raw tx; mempool timing correlation; for Ironwood, IP ↔ turnstile amount joins. Not "
        "assumed: decrypting note plaintexts from the wire, forging spends, or seeing wallet IP "
        "when submit is localhost/LAN Zebrad only.",
    )
    add_body(
        doc,
        "Defender goal: prevent “this IP submitted this txid / turnstile.” Biggest remote win is "
        "IP ↔ transaction submit. Sync-over-Nym with clearnet submit is a false sense of security.",
    )

    add_heading(doc, "4.2 Where the IP leak is (cases A1–A4)", 2)
    add_table(
        doc,
        ["Case", "Flow", "Stance"],
        [
            ("A1 Local Zebrad", "Wallet → LAN/localhost sendraw", "Preferred default; mixnet optional"),
            ("A2 Remote Zebrad clearnet", "Wallet → public :8232", "Must use Nym/Tor; refuse clearnet"),
            ("A3 Remote LWD submit", "SendTransaction to public LWD", "Biggest win: all txs over mixnet"),
            ("A4 dVPN sync + clearnet submit", "Sync private, submit clear", "Reject as “Nym done”"),
        ],
    )

    add_heading(doc, "4.3 Nym modes (B1–B4)", 2)
    add_bullet(
        doc,
        "B1 smolmix (mixnet): primary for submit / small RPC — sendraw, migrate-broadcast (#147).",
    )
    add_bullet(
        doc,
        "B2 smol-dvpn: bulk compact sync / censorship (#146); does not replace mixnet on submit.",
    )
    add_bullet(doc, "B3 mix-fetch: browser/extension later.")
    add_bullet(
        doc,
        "B4 system NymVPN + attestation: bridge only — do not claim “Nym integrated” from a checkbox.",
    )

    add_heading(doc, "4.4 Hybrid architecture", 2)
    add_formula(
        doc,
        "remote sync → smol-dvpn → LWD;  remote submit → smolmix → Zebrad/LWD;  "
        "always → baseline hygiene;  local Zebrad → direct JSON-RPC",
    )

    add_heading(doc, "4.5 Baseline hygiene", 2)
    add_table(
        doc,
        ["ID", "Mechanism", "Default"],
        [
            ("H1", "Checkpoint spacing (scan start)", "S = 256 blocks"),
            ("H2", "Max overlap rewind", "R_max = 128 blocks"),
            ("H3", "Broadcast delay", "D ~ U[30, 300] s"),
            ("H4", "Tip-sync guard", "T_guard = 120 s"),
        ],
    )
    add_formula(doc, "h* = S · floor(max(0, h0 − r) / S),  r ∈ [0, R_max],  h* ≤ h0")
    add_formula(doc, "proceed if t_now − t_sync ≥ T_guard")

    add_heading(doc, "4.6 Priority 1 modes and surfaces", 2)
    add_body(
        doc,
        "MigrationNetworkPrivacyMode: LocalNode (preferred), Tor/I2P detect, NymMixnetBroadcast "
        "(opt-in smolmix), attestation/force-clearnet as discouraged exception. CLI and desktop "
        "share ZebraClient::broadcast_transaction ± nym_mixnet_broadcast; api-server IP is the "
        "submit IP unless wrapped; mobile follows once FFI owns a session.",
    )

    add_heading(doc, "4.7 Evidence bar", 2)
    add_table(
        doc,
        ["ID", "Item", "Status"],
        [
            ("H1–H4", "Baseline hygiene", "Landed"),
            ("D2a", "smolmix IP relocate", "PASS (re-run 2026-07-26)"),
            ("D2c", "Opt-in mixnet broadcast", "Wired"),
            ("D2c-live", "Live remote sendraw over mixnet", "Open"),
            ("D2d", "Ironwood same egress", "Wired"),
            ("D1", "smol-dvpn LWD sync", "Spike"),
        ],
    )
    add_body(
        doc,
        "Do not claim product “Nym support” until D2b + D2c-live (or equivalent) PASS for remote "
        "submit. Local-node operators are already in the strong default without mixnet. Tracking: "
        "NYM_IP_PRIVACY_CASE_BREAKDOWN.md; issues #146, #147.",
    )

    add_heading(doc, "4.8 What Nym does not claim", 2)
    add_body(
        doc,
        "Nym/Tor do not erase KYC history, defeat subset-sum on revealed turnstile amounts, "
        "protect a compromised device, or replace ZIP 318 cohort privacy. Orchard/Ironwood hide "
        "note plaintexts; mixnets hide who talked to the network.",
    )
    doc.add_page_break()


def section_zns(doc):
    add_heading(doc, "5. Zcash Names (ZNS)", 1)
    add_body(
        doc,
        "Zcash Names (ZNS) maps names such as tacostand.zcash to an Orchard/Ironwood unified "
        "address (UA). Nozy does not operate the registry: we resolve, link to a Business "
        "profile, and later optionally claim once signing is clear. Protocol: zcashme/ZNS, "
        "zcashnames.com, zcashname-sdk. Backlog: issue #85, BUSINESS_ZEC_ZNS_TODO.md.",
    )
    add_heading(doc, "5.1 How ZNS works with Zcash", 2)
    add_bullet(
        doc,
        "Registration/update: shielded tx memo to registry UA; indexer verifies and serves JSON-RPC.",
    )
    add_bullet(
        doc,
        "Wallets resolve name → UA, then build a normal shielded send (identity layer, not a new pool).",
    )
    add_formula(
        doc,
        "name.zcash  →(indexer resolve)→  UA  →(ZIP-225 / Ironwood send)→  notes",
    )
    add_body(
        doc,
        "Nozy trust model: api-server proxies resolve; verify() indexer identity; TTL cache; "
        "configurable ZNS_MAINNET_URL / ZNS_TESTNET_URL. MVP claim path: external zcashnames.com "
        "+ link in-wallet — not holding admin keys; in-wallet CLAIM gated until sovereign/cosign clear.",
    )
    add_heading(doc, "5.2 Fit in Nozy (Personal / Business)", 2)
    add_table(
        doc,
        ["Concept", "Decision"],
        [
            ("Seed", "One BIP39 mnemonic"),
            ("Accounts", "Personal Orchard index 0; Business index 1"),
            ("Merchant identity", "Linked .zcash name → Business UA"),
        ],
    )
    add_heading(doc, "5.3 Future plans (roadmap)", 2)
    add_table(
        doc,
        ["Phase", "Scope"],
        [
            ("0", "Locked product decisions"),
            ("1", "Personal / Business profiles"),
            ("2", "Resolve + link; Send accepts name.zcash"),
            ("3 / 3b", "Sell mode QR POS; native invoices"),
            ("4–5", "ZIP-321 URIs; CSV ledger; UFVK disclosure"),
            ("6", "In-wallet claim/update (gated)"),
        ],
    )
    add_body(
        doc,
        "MVP: Business + Sell QR with linked name; customer pays shielded ZEC. ZNS does not "
        "replace Ironwood/Nym gates on migrate-broadcast.",
    )
    doc.add_page_break()


def section_phased_development(doc):
    add_heading(doc, "6. Phased Development and Releases", 1)
    add_heading(doc, "6.1 Phase table", 2)
    add_table(
        doc,
        ["Phase", "Name", "Shipped", "Gate"],
        [
            ("0", "Foundation", "HD wallet, Orchard scan, CLI", "Mainnet scan"),
            ("1", "Zebrad + LWD", "zeaking::lwd, compact SQLite", "Sync to tip"),
            ("2", "Mainnet send", "Witnesses, ZIP-225, broadcast", "Mainnet TXID"),
            ("3", "NU6.2 + pilot A1", "5-block expiry, ZIP-317 shape", "Fees + branch ID"),
            ("4", "Surfaces", "api-server, extension, desktop WIP", "Send/sync parity"),
            ("5", "Reliability (2026-06)", "BUG-2026-001–011", "Evidence PASS"),
            ("6", "Ironwood", "NU6.3, ZIP 318, desktop beta.2", "Migrate path"),
            ("7", "Network privacy", "Hygiene + Nym mixnet path", "Evidence before claims"),
            ("8+", "Observatory / business", "Pilot metrics, ZNS, Sell", "Per feature"),
        ],
    )
    add_heading(doc, "6.2 Release matrix", 2)
    add_table(
        doc,
        ["Surface", "Version line", "Notes"],
        [
            ("CLI / nozy", "v2.4.0 → v2.4.1 / v2.4.1.1", "Ironwood commands; mainnet CLI"),
            ("Prior CLI", "v2.3.0–v2.3.6.7", "Dynamic fee, NU6.2, lag guard"),
            ("Desktop", "1.0.0-beta.2", "Aligns with CLI v2.4.1.x"),
            ("Extension", "0.1.x beta", "Companion + WASM"),
            ("Mobile", "1.0.0 companion", "API + Zebrad you control"),
            ("api-server / zeaking", "0.1.x crates", "Local companion / sync"),
        ],
    )
    doc.add_page_break()


def section_challenges(doc):
    add_heading(doc, "7. Challenges and Responses", 1)
    add_heading(doc, "7.1 Zebrad integration", 2)
    add_body(
        doc,
        "Missing fee and witness RPCs are wallet problems on a Zebrad stack. Response: client "
        "ZIP-317 + local witnesses + treestate verification (ZEBRAD_SHIELDED_SEND_LIMIT.md).",
    )
    add_heading(doc, "7.2 Two clocks on shielded sends", 2)
    add_body(
        doc,
        "Build clock: witness → prove → sign → broadcast. Mempool clock starts after successful "
        "broadcast (Δ_exp = 5). Conflating them caused pre-broadcast −25 on slow VPS; fixed with "
        "late tip refresh and rebuild—not longer Δ_exp.",
    )
    add_heading(doc, "7.3 Ironwood migration UX vs privacy", 2)
    add_body(
        doc,
        "Schedule math is necessary but not sufficient. Users must understand turnstile amount "
        "disclosure and choose egress (local node vs mixnet).",
    )
    add_heading(doc, "7.4 Surface parity", 2)
    add_body(
        doc,
        "One fee/expiry/Ironwood policy across CLI, API, desktop, extension, and mobile remains "
        "the expensive ongoing gate.",
    )
    doc.add_page_break()


def section_tradeoffs(doc):
    add_heading(doc, "8. Trade-offs", 1)
    add_table(
        doc,
        ["Choice", "Alternative rejected", "Why"],
        [
            ("Zebrad + LWD", "Embed zcashd", "Wallet ≠ node"),
            ("Local witnesses", "Node witness RPC", "Not available"),
            ("Δ_exp = 5", "15", "Speed-up UX; rebuild instead"),
            ("F = 4 F_conv", "User-tunable fee UI", "Pilot simplicity"),
            ("ZIP 318 schedule", "One-shot migrate-all", "Amount/timing privacy"),
            ("Hygiene + optional Nym", "Claim anonymous by default", "Evidence-gated claims"),
            ("Ironwood after NU6.3", "Stay Orchard-only", "Pool sealed"),
        ],
    )
    doc.add_page_break()


def section_security(doc):
    add_heading(doc, "9. Security and Privacy Considerations", 1)
    add_bullet(doc, "High-impact code: keys, seeds, addresses, RPC URLs—no careless logging.")
    add_bullet(doc, "Shielded-first product policy (no accidental transparent sends).")
    add_bullet(doc, "Ironwood migration warnings for amount disclosure.")
    add_bullet(
        doc,
        "Nym/mixnet is opt-in / evidence-gated; attestation checkboxes are not proof of egress.",
    )
    add_bullet(
        doc,
        "No third-party audit claimed here; responsible disclosure per CONTRIBUTING.md / SECURITY.md.",
    )
    add_body(
        doc,
        "Note privacy follows Orchard/Ironwood ZK; network metadata is a separate layer (§4).",
    )
    doc.add_page_break()


def section_mainnet_evidence(doc):
    add_heading(doc, "10. Mainnet Field Evidence (June 2026)", 1)
    add_body(
        doc,
        "Operator stack: Windows host, Zebrad in WSL, nozy CLI release build. Dust amounts "
        "(0.0001 ZEC) for regression.",
    )
    add_body(doc, "Successful send TXIDs:")
    add_bullet(
        doc,
        "5a03fbd19547f9499182d78c88791eeb4eaab32e5d158b69ec8ffdc6068d2612",
    )
    add_bullet(
        doc,
        "902cf006efdeef3f15fed4312f8a15fcb1162f52495098c3bffb4acbe3cde4e5",
    )
    add_body(
        doc,
        "Observed T_send ~ 200 s when L ≤ 50. Proving warm-up ~ 2.1 s cold. Ironwood / Nym "
        "evidence continues under docs/reference/evidence/ (2026-07); see §4.3 for PASS vs open.",
    )
    doc.add_page_break()


def section_dynamic_fee_pilot(doc):
    add_heading(doc, "11. Dynamic-Fee Pilot Alignment", 1)
    add_table(
        doc,
        ["Pilot feature", "Implementation", "Lesson"],
        [
            ("Standard fee", "ZIP-317 in fee_policy.rs", "Zebrad has no fee RPC"),
            ("Priority ×4", "All Nozy send surfaces", "Speed-up after Expired"),
            ("Short expiry", "h_expiry = h_tip + 1 + 5", "Keep 5; rebuild for slow proves"),
            ("Speed-up", "New tx at priority fee", "Not rebroadcast of expired bytes"),
        ],
    )
    doc.add_page_break()


def section_lessons(doc):
    add_heading(doc, "12. Lessons Learned", 1)
    lessons = [
        "Two clocks: build-time expiry vs mempool expiry—fix the first without lengthening the second.",
        "Wallet ≠ node: witnesses and fees are wallet duties on Zebrad.",
        "Sync-before-send: L_max = 50 prevents pathological catch-up.",
        "Keep pilot knobs stable; improve runtime (rebuild, warm prove) first.",
        "Operator stacks (WSL Zebrad + Windows CLI) are first-class test targets.",
        "Cache migrations (NoteIndex v2) bite if any path still uses legacy parsers.",
        "Surface parity is expensive and mandatory.",
        "Ironwood migration is cryptography plus schedule math plus network metadata.",
        "Do not overclaim Nym: ship hygiene, wire mixnet, publish evidence.",
        "TXIDs and measured timings outperform narrative alone.",
    ]
    for i, lesson in enumerate(lessons, 1):
        add_body(doc, f"{i}. {lesson}")
    doc.add_page_break()


def section_conclusion(doc):
    add_heading(doc, "13. Conclusion", 1)
    add_body(
        doc,
        "NozyWallet shows a shielded-first wallet can run on Zebrad + lightwalletd with "
        "client-side ZIP-317 fees, short pilot expiry, and local witnesses—without zcashd and "
        "without stretching mempool expiry to hide slow proves. Ironwood extends that story to "
        "NU6.3 supply soundness and ZIP 318 migration. Nym-oriented broadcast and hygiene "
        "address the IP↔submit gap that ZK notes never covered. Releases through v2.4.1.x and "
        "desktop beta.2 carry this into operator hands.",
    )
    add_body(
        doc,
        "Continued work: Ironwood migrate hardening, honest Nym evidence for remote nodes, "
        "surface parity, pilot observatory metrics, and product features (ZNS, Sell) on a pure "
        "Zcash foundation.",
    )

    add_heading(doc, "References", 1)
    for r in [
        "LEONINE-DAO/Nozy-wallet — github.com/LEONINE-DAO/Nozy-wallet",
        "nozywallet.com",
        "docs/reference/MAINNET_SEND_READINESS_EVIDENCE.md",
        "docs/reference/PILOT_EXPIRY_PROVING_LATENCY.md",
        "docs/reference/IRONWOOD_WALLET_READINESS.md",
        "docs/reference/IRONWOOD_PRIVACY_EXPECTATIONS_ARTICLE.md",
        "docs/reference/NYM_IP_PRIVACY_CASE_BREAKDOWN.md",
        "docs/reference/SAFE_MIGRATION_NETWORK_PRIVACY_FORUM_POST.md",
        "ZEBRAD_SHIELDED_SEND_LIMIT.md",
        "book/src/features/ironwood.md",
        "Zebra — github.com/ZcashFoundation/zebra",
        "lightwalletd — github.com/zcash/lightwalletd",
        "ZIP-316, ZIP-317, ZIP-225; draft ZIP 318 (zips PR #1317)",
        "Shielded Labs Ironwood — shieldedlabs.net/ironwood/",
        "Nym Zcash guidance — zcash-sdk.nym.com",
    ]:
        add_bullet(doc, r)

    add_heading(doc, "Appendix A: Bug registry summary (2026-06)", 1)
    add_table(
        doc,
        ["ID", "Summary", "Status"],
        [
            ("BUG-2026-001", "Send rescanned ~50k blocks", "Fixed"),
            ("BUG-2026-002", "History empty despite balance", "Fixed"),
            ("BUG-2026-011", "Pre-broadcast expiry −25 on slow VPS", "Fixed"),
            ("—", "Witness lag guard, warm prove", "Fixed"),
            ("—", "NoteIndex v2 mark-spent", "Fixed"),
        ],
    )

    add_heading(doc, "Appendix B: Glossary", 1)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ("Anchor", "Commitment-tree root at a block height"),
            ("Ironwood", "NU6.3 shielded pool (corrected circuit)"),
            ("nExpiryHeight", "ZIP-225 last mineable height"),
            ("Pilot expiry", "Δ_exp = 5 after mempool build height"),
            ("Turnstile", "Accounted Orchard→Ironwood value exit"),
            ("Witness", "Merkle inclusion path for a note commitment"),
            ("ZIP 318", "Draft migration scheduling (buckets, K_max)"),
            ("Mixnet broadcast", "Opt-in sendraw via Nym smolmix egress"),
        ],
    )

    add_heading(doc, "Appendix C: Constant cheat-sheet", 1)
    add_table(
        doc,
        ["Name", "Value", "Source"],
        [
            ("f_m", "5,000 zat", "MARGINAL_FEE_ZATOSHIS"),
            ("g", "2", "GRACE_ACTIONS"),
            ("m", "4", "PRIORITY_MULTIPLIER"),
            ("Δ_exp", "5", "PILOT_EXPIRY_DELTA_BLOCKS"),
            ("L_max", "50 blocks", "Send-readiness"),
            ("B", "256 blocks", "ZIP318_ANCHOR_BUCKET_INTERVAL_BLOCKS"),
            ("K_max", "4", "ZIP318_DEFAULT_K_MAX"),
            ("r_min", "10^5 zat", "ZOOKO_RESIDUAL_ABANDON_ZAT"),
            ("H_act", "3,428,143", "NU6.3 mainnet"),
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of White Paper —")
    run.italic = True
    run.font.size = Pt(10)


def export_markdown(doc):
    """Plain-text markdown mirror for editors that cannot open .docx."""
    lines = [
        "# NozyWallet White Paper",
        "",
        "![NozyWallet logo](NozyWallet_Whitepaper_logo.png)",
        "",
    ]
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            lines.append("")
            continue
        style = (p.style.name or "") if p.style else ""
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif "List Bullet" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)
        lines.append("")
    OUTPUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def main(export_md: bool = False):
    doc = Document()
    set_margins(doc.sections[0])
    cover_page(doc)
    section_executive_summary(doc)
    section_architecture_decisions(doc)
    section_math(doc)
    section_ironwood(doc)
    section_nym(doc)
    section_zns(doc)
    section_phased_development(doc)
    section_challenges(doc)
    section_tradeoffs(doc)
    section_security(doc)
    section_mainnet_evidence(doc)
    section_dynamic_fee_pilot(doc)
    section_lessons(doc)
    section_conclusion(doc)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    if export_md:
        export_markdown(doc)
        print(f"Saved: {OUTPUT_MD}")
    else:
        print(
            "Note: docs/reference/NozyWallet_Whitepaper.md is hand-authored (LaTeX math). "
            "Pass --export-md to overwrite it from this Word build."
        )


if __name__ == "__main__":
    import sys

    main(export_md="--export-md" in sys.argv)
